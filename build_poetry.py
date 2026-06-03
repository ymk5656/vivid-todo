import os, re
import olefile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

UPLOAD_DIR = "/root/.claude/uploads/3eab84f8-b251-4dc6-b517-008b622cacb2/"
ZIP_DIR = os.path.join(UPLOAD_DIR, "zip_extracted/2025 #Uae40#Uc5f0#Ubbfc #Uc2dc #Ubaa8#Uc74c/")

PARTICLES = set("이가은는을를의와과도")
numbered_re = re.compile(r'^\d+\.\s+(.+)')

def decode_zip_fn(name):
    base = re.sub(r'\.(docx|hwp)$', '', name)
    return re.sub(r'#U([0-9a-fA-F]{4})', lambda m: chr(int(m.group(1),16)), base)

def split_para(t):
    return [l.strip() for l in t.split('\n') if l.strip()]

def content_key(path):
    try:
        doc = Document(path)
        lines = []
        for p in doc.paragraphs:
            for l in p.text.strip().split('\n'):
                if l.strip(): lines.append(l.strip())
                if len(lines) >= 3: return '||'.join(lines)
        return '||'.join(lines)
    except: return ''

def hwp_key(path):
    try:
        ole = olefile.OleFileIO(path)
        if not ole.exists('PrvText'): return ''
        d = ole.openstream('PrvText').read()
        t = d.decode('utf-16le',errors='ignore').replace('\r','\n').strip()
        lines = [l.strip() for l in t.split('\n') if l.strip()]
        return '||'.join(lines[:3])
    except: return ''

def extract_upload_title(fname, raw_paras):
    if not raw_paras: return '', []

    first_para = raw_paras[0]
    has_newline = '\n' in first_para

    name_part = re.sub(r'^[0-9a-f]{8}-', '', fname)
    name_part = re.sub(r'\.(docx|hwp)$', '', name_part)
    m = re.match(r'^(_+)(\d*)$', name_part)
    n = len(m.group(1)) if m else None
    suffix = m.group(2) if m else ''

    def try_rule2(line):
        if n is None: return None
        w = re.match(r'^([가-힣a-zA-Z]+)(?=[,\s])', line)
        if w and len(w.group(1)) == n:
            return w.group(1) + suffix
        return None

    def all_content(paras):
        cl = []
        for para in paras:
            cl.extend(split_para(para)); cl.append('')
        return cl

    if has_newline:
        first_line = first_para.split('\n')[0].strip()
        t2 = try_rule2(first_line)
        if t2:
            return t2, all_content(raw_paras)
        if n and len(first_line) == n + 1 and first_line[-1] in PARTICLES:
            return first_line[:-1] + suffix, all_content(raw_paras)
        return first_line + suffix, all_content(raw_paras)

    t2 = try_rule2(first_para)
    if t2:
        return t2, all_content(raw_paras)

    # Long first para with sentence breaks: use first sentence as title
    if n and len(first_para) > 30 and '. ' in first_para:
        parts = first_para.split('. ')
        first_sentence = parts[0].strip()
        remaining = [s.rstrip('.').strip() for s in parts[1:] if s.strip()]
        content_lines = remaining[:]
        content_lines.append('')
        for para in raw_paras[1:]:
            content_lines.extend(split_para(para)); content_lines.append('')
        return first_sentence, content_lines

    return first_para.strip(), all_content(raw_paras[1:])

def clean_title(title):
    title = ''.join(c for c in title if c.isprintable() or c == ' ').strip()
    title = title.rstrip('.')
    return title.strip()

def read_docx_all(path):
    try:
        doc = Document(path)
        return [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    except: return []

def read_hwp_lines(path):
    try:
        ole = olefile.OleFileIO(path)
        if not ole.exists('PrvText'): return []
        d = ole.openstream('PrvText').read()
        t = d.decode('utf-16le',errors='ignore').replace('\r','\n').strip()
        return [l.strip() for l in t.split('\n') if l.strip()]
    except: return []

def has_numbered_sections(raw_paras):
    if not raw_paras: return False
    first_line = raw_paras[0].split('\n')[0].strip()
    return bool(numbered_re.match(first_line))

def extract_numbered_poems(raw_paras):
    poems_out = []
    current_title = None
    current_lines = []

    for para in raw_paras:
        lines = para.split('\n')
        first_line = lines[0].strip()
        m = numbered_re.match(first_line)
        if m:
            if current_title:
                poems_out.append((current_title, current_lines))
            current_title = m.group(1).strip()
            rest = [l.strip() for l in lines[1:] if l.strip()]
            current_lines = rest + ['']
        else:
            if current_title:
                body = [l.strip() for l in lines if l.strip()]
                current_lines.extend(body)
                current_lines.append('')

    if current_title:
        poems_out.append((current_title, current_lines))

    return poems_out

# ── 수집 ─────────────────────────────────────────────────────────
poems = {}
zip_keys = {}

for f in sorted(os.listdir(ZIP_DIR)):
    path = os.path.join(ZIP_DIR, f)
    zt = decode_zip_fn(f)
    if f.endswith('.docx'):
        key = content_key(path)
        raw = read_docx_all(path)
        cl = []
        for para in raw: cl.extend(split_para(para)); cl.append('')
        if key: zip_keys[key] = zt
        zt_clean = clean_title(zt)
        if zt_clean and zt_clean not in poems:
            poems[zt_clean] = cl
    elif f.endswith('.hwp'):
        key = hwp_key(path)
        hl = read_hwp_lines(path)
        content = hl[1:] if hl and hl[0] == zt else hl
        if key: zip_keys[key] = zt
        zt_clean = clean_title(zt)
        if zt_clean and zt_clean not in poems:
            poems[zt_clean] = content
    else:
        continue

for f in sorted(os.listdir(UPLOAD_DIR)):
    if not f.endswith('.docx'): continue
    path = os.path.join(UPLOAD_DIR, f)
    key = content_key(path)
    if key in zip_keys: continue
    raw = read_docx_all(path)
    if not raw: continue
    if has_numbered_sections(raw):
        title, content = extract_upload_title(f, raw)
        title = clean_title(title)
        if title and title not in poems:
            cl = []
            for para in raw: cl.extend(split_para(para)); cl.append('')
            poems[title] = cl
        continue
    title, content = extract_upload_title(f, raw)
    title = clean_title(title)
    if title and title not in poems:
        poems[title] = content

for f in sorted(os.listdir(UPLOAD_DIR)):
    if not f.endswith('.hwp'): continue
    path = os.path.join(UPLOAD_DIR, f)
    key = hwp_key(path)
    if key in zip_keys: continue
    lines = read_hwp_lines(path)
    if not lines: continue
    title = clean_title(lines[0])
    content = lines[1:]
    if title and title not in poems:
        poems[title] = content

sorted_titles = sorted(poems.keys())
print(f"총 시 수: {len(sorted_titles)}편\n")
for i, t in enumerate(sorted_titles, 1):
    print(f"  {i:3d}. {t}")

# ── 디자인 빌드 ──────────────────────────────────────────────────
DARK = RGBColor(0x1a, 0x1a, 0x2e); MID = RGBColor(0x44, 0x44, 0x55)
LIGHT = RGBColor(0x99, 0x99, 0xaa); ACCT = RGBColor(0x6b, 0x4f, 0x3a); FT = "나눔명조"

def sf(run, sz, bold=False, col=None):
    run.font.name = FT; run.font.size = Pt(sz); run.font.bold = bold
    if col: run.font.color.rgb = col
    rPr = run._r.get_or_add_rPr(); rF = rPr.find(qn('w:rFonts'))
    if rF is None: rF = OxmlElement('w:rFonts'); rPr.insert(0, rF)
    rF.set(qn('w:eastAsia'), FT)

def sp(p, bef=0, aft=0, ln=None):
    p.paragraph_format.space_before = Pt(bef); p.paragraph_format.space_after = Pt(aft)
    if ln:
        pPr = p._p.get_or_add_pPr(); s = pPr.find(qn('w:spacing'))
        if s is None: s = OxmlElement('w:spacing'); pPr.append(s)
        s.set(qn('w:lineRule'), 'auto'); s.set(qn('w:line'), str(ln))

def border_bot(p, col="6b4f3a", sz=4):
    pPr = p._p.get_or_add_pPr(); bdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom'); b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), str(sz)); b.set(qn('w:space'), '6'); b.set(qn('w:color'), col)
    bdr.append(b); pPr.append(bdr)

def pgbr(p):
    pPr = p._p.get_or_add_pPr(); pb = OxmlElement('w:pageBreakBefore')
    pb.set(qn('w:val'), '1'); pPr.append(pb)

doc = Document()
for sec in doc.sections:
    sec.top_margin = Cm(3.2); sec.bottom_margin = Cm(3.2)
    sec.left_margin = Cm(4.0); sec.right_margin = Cm(4.0)

# 표지
for _ in range(7):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("✦"); sf(r, 16, col=ACCT); sp(p, aft=10)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("바람이 길을 묻는다"); sf(r, 30, bold=True, col=DARK); sp(p, bef=6, aft=10)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("시  집"); sf(r, 14, col=MID); sp(p, aft=60)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("김연민"); sf(r, 14, col=DARK); sp(p, aft=8)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("2025"); sf(r, 11, col=LIGHT)
doc.add_page_break()

# 목차
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("목  차"); sf(r, 20, bold=True, col=DARK)
border_bot(p, col="1a1a2e", sz=8); sp(p, bef=14, aft=22)
for i, title in enumerate(sorted_titles, 1):
    p = doc.add_paragraph()
    disp = title if len(title) <= 48 else title[:46] + "…"
    r1 = p.add_run(f"{i:3d}. "); sf(r1, 10.5, bold=True, col=ACCT)
    r2 = p.add_run(disp); sf(r2, 10.5, col=DARK); sp(p, aft=3)
doc.add_page_break()

# 본문
for idx, title in enumerate(sorted_titles):
    content = poems[title]
    p = doc.add_paragraph()
    if idx > 0: pgbr(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(title); sf(r, 17, bold=True, col=DARK)
    border_bot(p); sp(p, bef=28, aft=20, ln=276)
    stanza = []
    for line in content:
        if line.strip(): stanza.append(line)
        else:
            if stanza:
                pb = doc.add_paragraph('\n'.join(stanza))
                sf(pb.runs[0], 12.5, col=DARK); sp(pb, aft=14, ln=374); stanza = []
    if stanza:
        pb = doc.add_paragraph('\n'.join(stanza))
        sf(pb.runs[0], 12.5, col=DARK); sp(pb, aft=14, ln=374)

out = "/home/user/vivid-todo/시집_바람이길을묻는다.docx"
doc.save(out)
print(f"\n저장: {out}")
